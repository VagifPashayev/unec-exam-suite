import os
import json
import random
from tkinter import *
from tkinter import simpledialog, messagebox, ttk
from PIL import Image, ImageTk
from docx import Document
from docx.shared import Inches

DATA_FOLDER = "quiz_data"
METADATA_FILE = os.path.join(DATA_FOLDER, "metadata.json")
RESULT_BASENAME = "result"


class QuizApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Quiz")
        self.root.geometry("1200x800")
        self.root.minsize(800, 600)
        self.root.resizable(True, True)

        with open(METADATA_FILE, "r", encoding="utf-8") as f:
            self.metadata = json.load(f)

        self.score = 0
        self.wrong = 0
        self.streak = 0
        self.current = 0
        self.total = 0
        self.questions = []
        self.image_scale = 1.0
        self.result_path = None
        self.document = None

        self.ask_range()

    def _unique_result_path(self, qmin, qmax, count):
        for i in range(11):
            suffix = "" if i == 0 else f"_{i}"
            name = f"{RESULT_BASENAME} ({qmin}-{qmax} of {count}){suffix}.docx"
            if not os.path.exists(name):
                return name
        messagebox.showerror("Error", "Too many result files. Delete old ones and try again.")
        self.root.destroy()
        raise SystemExit

    def ask_range(self):
        available = sorted(int(k[1:]) for k in self.metadata if k.startswith("q"))
        lo, hi = min(available), max(available)

        qmin = simpledialog.askinteger("Start", f"From question: (available: {lo}–{hi})", minvalue=1, parent=self.root)
        if qmin is None:
            self.root.destroy()
            return
        qmax = simpledialog.askinteger("Start", f"To question: (available: {lo}–{hi})", minvalue=qmin, parent=self.root)
        if qmax is None:
            self.root.destroy()
            return
        count = simpledialog.askinteger("Start", "How many questions:", minvalue=1, parent=self.root)
        if count is None:
            self.root.destroy()
            return

        pool = [f"q{i}" for i in range(qmin, qmax + 1) if f"q{i}" in self.metadata]
        if not pool:
            messagebox.showerror("Error", f"No questions in range {qmin}–{qmax}")
            self.root.destroy()
            return

        count = min(count, len(pool))
        self.total = count
        self.result_path = self._unique_result_path(qmin, qmax, count)
        self.document = Document()
        self.document.add_heading("Quiz Results", level=1)

        self.setup_gui()
        self.questions = random.sample(pool, count)
        self.load_question()

    def setup_gui(self):
        self.frame = Frame(self.root, bg="white")
        self.frame.pack(fill=BOTH, expand=True)

        self.progress_var = IntVar()
        self.progress = ttk.Progressbar(self.frame, maximum=100, variable=self.progress_var)
        self.progress.pack(pady=5, fill=X, padx=20)

        self.stats_label = Label(self.frame, text="", bg="white", font=("Arial", 12))
        self.stats_label.pack(pady=5)

        self.canvas = Canvas(self.frame, bg="white")
        self.scroll_y = Scrollbar(self.frame, orient="vertical", command=self.canvas.yview)
        self.scroll_x = Scrollbar(self.frame, orient="horizontal", command=self.canvas.xview)
        self.scroll_y.pack(side=RIGHT, fill=Y)
        self.scroll_x.pack(side=BOTTOM, fill=X)
        self.canvas.pack(side=LEFT, fill=BOTH, expand=True)
        self.canvas.configure(yscrollcommand=self.scroll_y.set, xscrollcommand=self.scroll_x.set)

        self.inner_frame = Frame(self.canvas, bg="white")
        self.canvas_window = self.canvas.create_window((0, 0), window=self.inner_frame, anchor="n")
        self.inner_frame.bind("<Configure>", self._on_inner_resize)

        self.image_label = Label(self.inner_frame, bg="white")
        self.image_label.pack(pady=10)

        self.root.bind("+", lambda e: self.zoom_in())
        self.root.bind("-", lambda e: self.zoom_out())

        self.option_buttons = []
        self.option_frames = []

    def _on_inner_resize(self, event):
        self.canvas.update_idletasks()
        cw = self.canvas.winfo_width()
        iw = self.inner_frame.winfo_reqwidth()
        x = max((cw - iw) // 2, 0)
        self.canvas.coords(self.canvas_window, x, 0)
        self.canvas.configure(scrollregion=self.canvas.bbox("all"))

    def zoom_in(self):
        self.image_scale *= 1.2
        self.load_question(refresh=True)

    def zoom_out(self):
        self.image_scale *= 0.8
        self.load_question(refresh=True)

    def load_question(self, refresh=False):
        if self.current >= self.total:
            self.finish()
            return

        if refresh:
            for btn in self.option_buttons:
                btn.destroy()
            for frm in self.option_frames:
                frm.destroy()
            self.option_buttons.clear()
            self.option_frames.clear()
        else:
            for w in self.inner_frame.winfo_children():
                if w is not self.image_label:
                    w.destroy()
            self.option_buttons.clear()
            self.option_frames.clear()

        pct = int(self.current / self.total * 100)
        self.progress_var.set(pct)
        self.stats_label.config(
            text=f"Correct: {self.score}  |  Wrong: {self.wrong}  |  Streak: {self.streak}  |  {self.current + 1}/{self.total}"
        )

        qkey = self.questions[self.current]
        qdata = self.metadata[qkey]

        img = Image.open(os.path.join(DATA_FOLDER, qdata["question"]))
        w, h = img.size
        img = img.resize((int(w * self.image_scale), int(h * self.image_scale)))
        self.q_photo = ImageTk.PhotoImage(img)
        self.image_label.config(image=self.q_photo)

        options = list(qdata["options"].items())
        random.shuffle(options)
        self.current_options = options
        self.correct_answer = qdata["correct"]
        self.correct_index = next(i for i, (lbl, _) in enumerate(options) if lbl == self.correct_answer)
        self.correct_letter = chr(65 + self.correct_index)

        for i, (label, img_file) in enumerate(options):
            frm = Frame(self.inner_frame, bg="white")
            frm.pack(pady=5)
            photo = ImageTk.PhotoImage(Image.open(os.path.join(DATA_FOLDER, img_file)))
            btn = Button(frm, text=f"{chr(65 + i)})", width=5,
                         command=lambda l=label, idx=i: self.check_answer(l, idx))
            lbl_img = Label(frm, image=photo, bg="white")
            lbl_img.image = photo
            btn.pack(side=LEFT)
            lbl_img.pack(side=LEFT)
            self.option_buttons.append(btn)
            self.option_frames.append(frm)

    def check_answer(self, chosen, idx):
        qkey = self.questions[self.current]
        chosen_letter = chr(65 + idx)
        displayed = [(chr(65 + i), lbl, img) for i, (lbl, img) in enumerate(self.current_options)]

        if chosen == self.correct_answer:
            self.score += 1
            self.streak += 1
        else:
            self.wrong += 1
            self.streak = 0
            messagebox.showerror("Wrong", f"Correct answer: {self.correct_letter}")
            self._save_wrong(qkey, chosen_letter, self.correct_letter, displayed)

        self.current += 1
        self.load_question()

    def _save_wrong(self, qkey, chosen, correct, displayed):
        self.document.add_heading(f"Question: {qkey}", level=2)
        self.document.add_picture(
            os.path.join(DATA_FOLDER, self.metadata[qkey]["question"]), width=Inches(3.5)
        )
        for letter, _, img_file in displayed:
            self.document.add_paragraph(f"{letter})")
            self.document.add_picture(os.path.join(DATA_FOLDER, img_file))
        self.document.add_paragraph(f"Your answer: {chosen}")
        self.document.add_paragraph(f"Correct answer: {correct}")

    def finish(self):
        self.document.add_paragraph(
            f"\n\nCorrect: {self.score}/{self.total}\nWrong: {self.wrong}\nBest streak: {self.streak}"
        )
        self.document.save(self.result_path)
        messagebox.showinfo(
            "Done",
            f"Quiz finished!\nCorrect: {self.score}/{self.total}\nResult saved to: {self.result_path}",
        )
        self.root.destroy()


if __name__ == "__main__":
    root = Tk()
    app = QuizApp(root)
    root.mainloop()
