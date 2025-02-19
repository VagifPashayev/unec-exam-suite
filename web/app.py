import os
import json
import random
from datetime import datetime
from docx import Document
from docx.shared import Inches
from flask import Flask, render_template, request, session, redirect, url_for, send_file

app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", "change-me-in-production")

DATA_FOLDER = os.path.join(os.path.dirname(__file__), "static", "quiz_data")
METADATA_FILE = os.path.join(DATA_FOLDER, "metadata.json")
RESULT_FOLDER = os.path.join(os.path.dirname(__file__), "results")
LETTERS = ["A", "B", "C", "D", "E"]

with open(METADATA_FILE, encoding="utf-8") as f:
    METADATA = json.load(f)

os.makedirs(RESULT_FOLDER, exist_ok=True)


def _available_range():
    keys = sorted(int(k[1:]) for k in METADATA if k.startswith("q"))
    return min(keys), max(keys)


@app.route("/", methods=["GET", "POST"])
def index():
    min_q, max_q = _available_range()

    if request.method == "POST":
        custom = request.form.get("custom", "").strip()
        if custom:
            try:
                numbers = [int(x) for x in custom.split()]
                questions = [f"q{n}" for n in numbers if f"q{n}" in METADATA]
            except ValueError:
                questions = []
        else:
            qmin = int(request.form["qmin"])
            qmax = int(request.form["qmax"])
            count = int(request.form["count"])
            pool = [f"q{i}" for i in range(qmin, qmax + 1) if f"q{i}" in METADATA]
            questions = random.sample(pool, min(count, len(pool)))

        if not questions:
            return render_template("index.html", min_q=min_q, max_q=max_q, error="No questions found.")

        session.update({
            "questions": questions,
            "count": len(questions),
            "current": 0,
            "score": 0,
            "wrong": 0,
            "streak": 0,
            "review": None,
            "result_file": f"quiz_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
        })

        doc = Document()
        doc.add_heading("Quiz Results", level=1)
        doc.save(os.path.join(RESULT_FOLDER, session["result_file"]))

        return redirect(url_for("quiz"))

    return render_template("index.html", min_q=min_q, max_q=max_q)


@app.route("/quiz")
def quiz():
    current = session.get("current", 0)
    questions = session.get("questions", [])

    if current >= len(questions):
        return redirect(url_for("finish"))

    qkey = questions[current]
    qdata = METADATA[qkey]
    options = list(qdata["options"].items())
    random.shuffle(options)

    session["correct_label"] = qdata["correct"]
    session["options"] = options
    session["original_options"] = list(qdata["options"].items())

    return render_template(
        "quiz.html",
        question_id=current + 1,
        total=len(questions),
        stats={"score": session["score"], "wrong": session["wrong"], "streak": session["streak"]},
        qimage=qdata["question"],
        options=options,
    )


@app.route("/answer", methods=["POST"])
def answer():
    choice = request.form["choice"]
    correct = session["correct_label"]
    options = session["options"]
    original_options = session["original_options"]
    current = session["current"]
    qkey = session["questions"][current]
    result_path = os.path.join(RESULT_FOLDER, session["result_file"])

    if choice == correct:
        session["score"] += 1
        session["streak"] += 1
        session["current"] += 1
        return redirect(url_for("quiz"))

    session["wrong"] += 1
    session["streak"] = 0

    doc = Document(result_path)
    doc.add_heading(f"Question: {qkey}", level=2)
    doc.add_picture(os.path.join(DATA_FOLDER, METADATA[qkey]["question"]), width=Inches(4))
    for idx, (_, img) in enumerate(original_options):
        doc.add_paragraph(f"{LETTERS[idx]})")
        doc.add_picture(os.path.join(DATA_FOLDER, img), width=Inches(2.5))
    correct_letter = next((LETTERS[i] for i, (_, img) in enumerate(original_options) if img == correct), "?")
    chosen_letter = next((LETTERS[i] for i, (_, img) in enumerate(original_options) if img == choice), "?")
    doc.add_paragraph(f"Your answer: {chosen_letter}")
    doc.add_paragraph(f"Correct answer: {correct_letter}")
    doc.save(result_path)

    session["review"] = {
        "qkey": qkey,
        "question_img": METADATA[qkey]["question"],
        "options": options,
        "correct": correct,
        "chosen": choice,
        "correct_letter": next((LETTERS[i] for i, (lbl, _) in enumerate(options) if lbl == correct), "?"),
        "chosen_letter": next((LETTERS[i] for i, (lbl, _) in enumerate(options) if lbl == choice), "?"),
    }
    return redirect(url_for("review"))


@app.route("/review")
def review():
    data = session.get("review")
    if not data:
        return redirect(url_for("quiz"))
    return render_template(
        "review.html",
        question_img=data["question_img"],
        options=data["options"],
        correct=data["correct"],
        chosen=data["chosen"],
        correct_letter=data["correct_letter"],
        chosen_letter=data["chosen_letter"],
    )


@app.route("/next", methods=["POST"])
def next_question():
    session["current"] += 1
    session["review"] = None
    return redirect(url_for("quiz"))


@app.route("/finish")
def finish():
    result_path = os.path.join(RESULT_FOLDER, session["result_file"])
    doc = Document(result_path)

    score = session["score"]
    total = session["count"]
    wrong = session["wrong"]
    streak = session["streak"]

    wrong_nums = []
    for para in doc.paragraphs:
        if para.text.startswith("Question: q"):
            try:
                wrong_nums.append(para.text.split("q")[1])
            except IndexError:
                pass

    doc.add_paragraph(f"\nCorrect: {score}/{total}")
    doc.add_paragraph(f"Wrong: {wrong}")
    doc.add_paragraph(f"Best streak: {streak}")
    if wrong_nums:
        doc.add_paragraph("Wrong questions: " + " ".join(wrong_nums))
    doc.save(result_path)

    return render_template(
        "finish.html",
        score=score,
        total=total,
        streak=streak,
        file=session["result_file"],
    )


@app.route("/download/<filename>")
def download(filename):
    return send_file(os.path.join(RESULT_FOLDER, filename), as_attachment=True)


if __name__ == "__main__":
    app.run(debug=False, port=5000)
