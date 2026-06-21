import os
import sys
from datetime import datetime

from docx import Document

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "telegram_quiz"))

from quiz import build_result_document, prepare_display_options, safe_report_filename


class ReverseRng:
    @staticmethod
    def shuffle(values):
        values.reverse()


def test_prepare_display_options_keeps_answer_mapping_after_shuffle():
    question = {
        "number": 7,
        "question": "Question",
        "options": [
            {"id": "a", "text": "Wrong"},
            {"id": "b", "text": "Correct"},
            {"id": "c", "text": "Other"},
        ],
        "answer": "b",
    }
    displayed, correct_letter = prepare_display_options(question, rng=ReverseRng())
    assert [item["text"] for item in displayed] == ["Other", "Correct", "Wrong"]
    assert correct_letter == "b"
    assert displayed[1]["correct"] is True


def test_result_document_contains_displayed_selected_and_correct_answers(tmp_path):
    incorrect = [
        {
            "number": 42,
            "question": "Clinical question?",
            "options": [
                {"letter": "a", "text": "First", "correct": False},
                {"letter": "b", "text": "Second", "correct": True},
            ],
            "selected_letter": "a",
            "selected_text": "First",
            "correct_letter": "b",
            "correct_text": "Second",
        }
    ]
    document = build_result_document(
        test_name="Ginekologiya",
        score=0,
        total=1,
        max_streak=0,
        incorrect=incorrect,
        language="en",
        timestamp=datetime(2026, 6, 21, 12, 0),
    )
    output = tmp_path / "result.docx"
    document.save(output)
    loaded = Document(output)
    text = "\n".join(paragraph.text for paragraph in loaded.paragraphs)
    assert "Question 42" in text
    assert "Your answer: A) First" in text
    assert "Correct answer: B) Second" in text
    assert any("Correct answers" in cell.text for table in loaded.tables for row in table.rows for cell in row.cells)


def test_safe_report_filename_removes_windows_reserved_characters():
    assert safe_report_filename('bad:name/test', 1, 20) == "bad_name_test (1-20).docx"
