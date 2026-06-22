import json
import os
import sys
from pathlib import Path

import pytest
from docx import Document

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "telegram_quiz"))

from file_service import delete_quiz_file, install_quiz_file, validate_quiz_file
from loader import QuizFormatError
from storage import BotStorage


def _quiz_docx(
    path: Path, *, duplicate_number: bool = False, table: bool = False
) -> None:
    document = Document()
    document.add_paragraph("1. First question?")
    document.add_paragraph("a) Wrong")
    document.add_paragraph("b) Correct [[CORRECT]]")
    number = 1 if duplicate_number else 2
    document.add_paragraph(f"{number}. Second question?")
    document.add_paragraph("a) Yes [[CORRECT]]")
    document.add_paragraph("b) No")
    if table:
        document.add_table(rows=1, cols=1).cell(0, 0).text = "hidden question"
    document.save(path)


def test_storage_roles_access_and_owner_protection(tmp_path):
    storage = BotStorage(tmp_path / "bot.db")
    storage.initialize(100)
    assert storage.is_owner(100)
    assert storage.is_admin(100)

    storage.upsert_identity(200, "student", "Student")
    storage.set_language(200, "ru")
    assert storage.status(200) == "new"
    assert storage.request_access(200) is True
    assert storage.request_access(200) is False
    assert storage.status(200) == "pending"
    assert storage.set_access(100, 200, True) is True
    assert storage.is_approved(200)

    assert storage.grant_admin(100, 200) is True
    assert storage.is_admin(200)
    assert storage.revoke_admin(100, 200) is True
    assert storage.role(200) == "user"
    assert storage.set_access(200, 100, False) is False
    assert storage.is_owner(100)


def test_storage_migrates_legacy_json_once(tmp_path):
    approved = tmp_path / "approved_users.json"
    pending = tmp_path / "pending_users.json"
    languages = tmp_path / "user_languages.json"
    approved.write_text(json.dumps([200]), encoding="utf-8")
    pending.write_text(json.dumps([300]), encoding="utf-8")
    languages.write_text(json.dumps({"200": "ru", "300": "az"}), encoding="utf-8")

    storage = BotStorage(tmp_path / "bot.db")
    storage.initialize(
        100,
        approved_file=approved,
        pending_file=pending,
        languages_file=languages,
    )
    assert storage.is_approved(200)
    assert storage.get_language(200) == "ru"
    assert storage.status(300) == "pending"
    assert storage.get_language(300) == "az"

    approved.write_text(json.dumps([400]), encoding="utf-8")
    storage.initialize(100, approved_file=approved)
    assert storage.status(400) is None


def test_install_validate_download_metadata_and_soft_delete(tmp_path):
    storage = BotStorage(tmp_path / "bot.db")
    storage.initialize(100)
    source = tmp_path / "source.docx"
    quiz_dir = tmp_path / "quizzes"
    trash_dir = tmp_path / "trash"
    _quiz_docx(source)

    report = validate_quiz_file(source)
    assert report.question_count == 2
    item, installed_report = install_quiz_file(
        source, "Clinical: test.docx", 100, storage, quiz_dir
    )
    assert item.filename == "Clinical_ test.docx"
    assert installed_report.sha256 == report.sha256
    assert (quiz_dir / item.filename).read_bytes() == source.read_bytes()

    deleted = delete_quiz_file(item.id, 100, storage, quiz_dir, trash_dir)
    assert deleted.id == item.id
    assert storage.get_quiz(item.id) is None
    assert not (quiz_dir / item.filename).exists()
    assert (trash_dir / f"{item.id}-{item.filename}").exists()


def test_validator_rejects_duplicate_numbers_and_table_content(tmp_path):
    duplicate = tmp_path / "duplicate.docx"
    _quiz_docx(duplicate, duplicate_number=True)
    with pytest.raises(QuizFormatError, match="duplicate question number"):
        validate_quiz_file(duplicate)

    table = tmp_path / "table.docx"
    _quiz_docx(table, table=True)
    with pytest.raises(QuizFormatError, match="tables"):
        validate_quiz_file(table)
