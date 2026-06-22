"""Safe quiz-bank installation, validation, and removal."""

from __future__ import annotations

import hashlib
import os
import re
import shutil
import tempfile
import zipfile
from dataclasses import dataclass
from html import escape
from pathlib import Path

from docx import Document

from loader import QuizFormatError, load_questions_from_docx
from storage import BotStorage, QuizFile


MAX_DOCX_BYTES = 20 * 1024 * 1024
MAX_UNCOMPRESSED_BYTES = 60 * 1024 * 1024
MAX_ARCHIVE_MEMBERS = 3000
MAX_QUESTIONS = 10000
MAX_OPTIONS = 26


@dataclass(frozen=True)
class ValidationReport:
    question_count: int
    sha256: str
    warnings: tuple[str, ...]


def _safe_filename(name: str) -> str:
    name = Path(name.replace("\\", "/")).name
    stem = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', "_", Path(name).stem).strip(" .")
    if not stem:
        stem = "quiz"
    return f"{stem[:100]}.docx"


def _validate_archive(path: Path) -> None:
    if path.suffix.lower() != ".docx":
        raise QuizFormatError("only .docx files are supported")
    size = path.stat().st_size
    if size <= 0 or size > MAX_DOCX_BYTES:
        raise QuizFormatError(
            f"file size must be between 1 byte and {MAX_DOCX_BYTES // 1024 // 1024} MB"
        )
    if not zipfile.is_zipfile(path):
        raise QuizFormatError("the file is not a valid DOCX archive")
    with zipfile.ZipFile(path) as archive:
        members = archive.infolist()
        if len(members) > MAX_ARCHIVE_MEMBERS:
            raise QuizFormatError("the DOCX archive contains too many entries")
        total = 0
        for member in members:
            member_path = Path(member.filename.replace("\\", "/"))
            if member_path.is_absolute() or ".." in member_path.parts:
                raise QuizFormatError("the DOCX archive contains an unsafe path")
            total += member.file_size
            if total > MAX_UNCOMPRESSED_BYTES:
                raise QuizFormatError("the DOCX archive is too large when unpacked")


def validate_quiz_file(path: str | Path) -> ValidationReport:
    path = Path(path)
    _validate_archive(path)
    try:
        document = Document(path)
    except Exception as error:
        raise QuizFormatError(f"DOCX cannot be opened: {error}") from error
    if any(
        cell.text.strip()
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    ):
        raise QuizFormatError(
            "tables with question content are not supported; use ordinary paragraphs"
        )

    questions = load_questions_from_docx(path)
    if not questions:
        raise QuizFormatError("no questions found")
    if len(questions) > MAX_QUESTIONS:
        raise QuizFormatError(f"more than {MAX_QUESTIONS} questions")

    numbers: set[int] = set()
    texts: set[str] = set()
    duplicate_texts = 0
    for question in questions:
        number = question["number"]
        if number in numbers:
            raise QuizFormatError(f"duplicate question number {number}")
        numbers.add(number)
        if not question["question"].strip():
            raise QuizFormatError(f"question {number} has empty text")
        if len(question["options"]) > MAX_OPTIONS:
            raise QuizFormatError(
                f"question {number} has more than {MAX_OPTIONS} options"
            )
        normalized = re.sub(r"\s+", " ", question["question"]).strip().casefold()
        if normalized in texts:
            duplicate_texts += 1
        texts.add(normalized)
        rendered = (
            escape(question["question"])
            + "\n"
            + "\n".join(
                f"{option['id']}) {escape(option['text'])}"
                for option in question["options"]
            )
        )
        if len(rendered) + 300 > 4096:
            raise QuizFormatError(f"question {number} exceeds Telegram's message limit")

    warnings: list[str] = []
    raw_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    if "[[CORRECT]]" not in raw_text and "(+)" in raw_text:
        warnings.append("legacy (+) answer markers are accepted; [[CORRECT]] is safer")
    if duplicate_texts:
        warnings.append(f"{duplicate_texts} duplicate question texts found")
    digest = hashlib.sha256(path.read_bytes()).hexdigest()
    return ValidationReport(len(questions), digest, tuple(warnings))


def initialize_quiz_files(
    storage: BotStorage, bundled_dir: Path, quiz_dir: Path
) -> None:
    quiz_dir.mkdir(parents=True, exist_ok=True)
    if bundled_dir.resolve() != quiz_dir.resolve() and not storage.list_quizzes(
        active_only=False
    ):
        for source in bundled_dir.glob("*.docx"):
            destination = quiz_dir / _safe_filename(source.name)
            if not destination.exists():
                shutil.copy2(source, destination)
    known = {item.filename for item in storage.list_quizzes(active_only=False)}
    for path in quiz_dir.glob("*.docx"):
        if path.name in known:
            continue
        report = validate_quiz_file(path)
        storage.register_quiz(
            filename=path.name,
            title=path.stem,
            sha256=report.sha256,
            question_count=report.question_count,
            uploaded_by=None,
        )


def install_quiz_file(
    source: str | Path,
    original_name: str,
    actor_id: int,
    storage: BotStorage,
    quiz_dir: Path,
) -> tuple[QuizFile, ValidationReport]:
    source = Path(source)
    filename = _safe_filename(original_name)
    destination = quiz_dir / filename
    if destination.exists():
        raise QuizFormatError(f"a file named {filename} already exists")
    report = validate_quiz_file(source)
    quiz_dir.mkdir(parents=True, exist_ok=True)
    handle, temporary_name = tempfile.mkstemp(
        prefix=".upload-", suffix=".docx", dir=quiz_dir
    )
    os.close(handle)
    temporary_path = Path(temporary_name)
    try:
        shutil.copyfile(source, temporary_path)
        os.replace(temporary_path, destination)
        try:
            item = storage.register_quiz(
                filename=filename,
                title=Path(filename).stem,
                sha256=report.sha256,
                question_count=report.question_count,
                uploaded_by=actor_id,
            )
        except Exception:
            destination.unlink(missing_ok=True)
            raise
    finally:
        temporary_path.unlink(missing_ok=True)
    return item, report


def delete_quiz_file(
    quiz_id: int, actor_id: int, storage: BotStorage, quiz_dir: Path, trash_dir: Path
) -> QuizFile | None:
    item = storage.get_quiz(quiz_id)
    if not item:
        return None
    source = quiz_dir / item.filename
    trash_dir.mkdir(parents=True, exist_ok=True)
    destination = trash_dir / f"{item.id}-{item.filename}"
    if source.exists():
        os.replace(source, destination)
    try:
        return storage.deactivate_quiz(actor_id, quiz_id)
    except Exception:
        if destination.exists():
            os.replace(destination, source)
        raise
