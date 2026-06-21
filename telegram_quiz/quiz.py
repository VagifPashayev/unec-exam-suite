"""Quiz presentation and deterministic Word result reporting."""

from __future__ import annotations

import os
import random
import re
import tempfile
from datetime import datetime
from html import escape
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from telegram import InlineKeyboardButton, InlineKeyboardMarkup

from i18n import t
from utils import build_progress_bar


DISPLAY_LABELS = "abcdefghijklmnopqrstuvwxyz"


def prepare_display_options(question: dict, rng=random) -> tuple[list[dict], str]:
    options = [dict(option) for option in question["options"]]
    rng.shuffle(options)
    displayed: list[dict] = []
    correct_letter = ""
    for index, option in enumerate(options):
        letter = DISPLAY_LABELS[index]
        item = {
            "letter": letter,
            "source_id": option["id"],
            "text": option["text"],
            "correct": option["id"] == question["answer"],
        }
        displayed.append(item)
        if item["correct"]:
            correct_letter = letter
    if not correct_letter:
        raise ValueError(f"question {question.get('number')} has no matching correct answer")
    return displayed, correct_letter


async def ask_question(update, context) -> None:
    index = context.user_data["index"]
    question = context.user_data["questions"][index]
    language = context.user_data.get("language", "en")
    displayed, correct_letter = prepare_display_options(question)

    context.user_data["displayed_options"] = displayed
    context.user_data["correct_letter"] = correct_letter

    total = len(context.user_data["questions"])
    correct = context.user_data["score"]
    wrong = index - correct
    bar = build_progress_bar(index, total)
    header = t(language, "question_progress", current=index + 1, total=total)
    text = (
        f"<b>{escape(header)}</b>  {bar}\n"
        f"✅ {correct} | ❌ {wrong} | 🔥 {context.user_data.get('streak', 0)}\n\n"
        f"<b>{escape(question['question'])}</b>\n\n"
    )
    for option in displayed:
        text += f"{option['letter']}) {escape(option['text'])}\n"

    if len(text) > 4096:
        raise ValueError(f"question {question.get('number')} exceeds Telegram's message limit")

    keyboard = [
        [
            InlineKeyboardButton(
                option["letter"].upper(),
                callback_data=f"answer:{index}:{option['letter']}",
            )
        ]
        for option in displayed
    ]
    keyboard.append(
        [InlineKeyboardButton(t(language, "main_menu"), callback_data="main_menu")]
    )
    await context.bot.send_message(
        chat_id=update.effective_chat.id,
        text=text,
        reply_markup=InlineKeyboardMarkup(keyboard),
        parse_mode="HTML",
    )


def _set_cell_width(cell, width_dxa: int) -> None:
    properties = cell._tc.get_or_add_tcPr()
    width = properties.find(qn("w:tcW"))
    if width is None:
        width = OxmlElement("w:tcW")
        properties.append(width)
    width.set(qn("w:w"), str(width_dxa))
    width.set(qn("w:type"), "dxa")


def _shade_cell(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def _add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def build_result_document(
    *,
    test_name: str,
    score: int,
    total: int,
    max_streak: int,
    incorrect: list[dict],
    language: str,
    timestamp: datetime | None = None,
) -> Document:
    """Build a compact report that keeps displayed and correct answers aligned."""

    now = timestamp or datetime.now()
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading = document.styles["Heading 1"]
    heading.font.name = "Calibri"
    heading.font.size = Pt(16)
    heading.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    heading.paragraph_format.space_before = Pt(18)
    heading.paragraph_format.space_after = Pt(10)

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run(t(language, "report_title"))
    title_run.bold = True
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)

    subtitle = document.add_paragraph(test_name)
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle.runs[0].font.size = Pt(12)
    subtitle.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    rows = [
        (t(language, "report_test"), test_name),
        (t(language, "report_date"), now.strftime("%Y-%m-%d %H:%M")),
        (t(language, "report_score"), f"{score} / {total}"),
        (t(language, "report_incorrect"), str(len(incorrect))),
        (t(language, "report_streak"), str(max_streak)),
    ]
    table = document.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row, (label, value) in zip(table.rows, rows):
        _set_cell_width(row.cells[0], 2700)
        _set_cell_width(row.cells[1], 6660)
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _shade_cell(row.cells[0], "E8EEF5")
        label_run = row.cells[0].paragraphs[0].add_run(label)
        label_run.bold = True
        row.cells[1].paragraphs[0].add_run(value)

    if not incorrect:
        paragraph = document.add_paragraph(t(language, "report_no_errors"))
        paragraph.paragraph_format.space_before = Pt(14)
    else:
        document.add_heading(t(language, "report_errors"), level=1)
        for item in incorrect:
            document.add_heading(
                t(language, "report_question", number=item["number"]), level=1
            )
            question = document.add_paragraph(item["question"])
            question.runs[0].bold = True

            for option in item["options"]:
                paragraph = document.add_paragraph(style="List Bullet")
                run = paragraph.add_run(
                    f"{option['letter'].upper()}) {option['text']}"
                )
                if option["letter"] == item["selected_letter"]:
                    run.font.color.rgb = RGBColor(0x9B, 0x1C, 0x1C)
                if option["letter"] == item["correct_letter"]:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0x1F, 0x5E, 0x35)

            selected = document.add_paragraph()
            selected_label = selected.add_run(f"{t(language, 'your_answer')}: ")
            selected_label.bold = True
            selected_label.font.color.rgb = RGBColor(0x9B, 0x1C, 0x1C)
            selected.add_run(
                f"{item['selected_letter'].upper()}) {item['selected_text']}"
            )

            correct = document.add_paragraph()
            correct_label = correct.add_run(f"{t(language, 'correct_answer')}: ")
            correct_label.bold = True
            correct_label.font.color.rgb = RGBColor(0x1F, 0x5E, 0x35)
            correct.add_run(
                f"{item['correct_letter'].upper()}) {item['correct_text']}"
            )

    footer = section.footer.paragraphs[0]
    _add_page_number(footer)
    return document


def safe_report_filename(test_name: str, start: int, end: int) -> str:
    safe_name = re.sub(r'[<>:"/\\|?*]+', "_", test_name).strip(" .") or "quiz"
    return f"{safe_name} ({start}-{end}).docx"


async def send_result_doc(update, context) -> None:
    questions = context.user_data.get("questions", [])
    score = context.user_data.get("score", 0)
    incorrect = context.user_data.get("incorrect_details", [])
    selected_file = context.user_data.get("selected_file", "quiz.docx")
    test_name = Path(selected_file).stem
    start = context.user_data.get("q_start", 1)
    end = context.user_data.get("q_end", start + max(len(questions) - 1, 0))
    language = context.user_data.get("language", "en")

    document = build_result_document(
        test_name=test_name,
        score=score,
        total=len(questions),
        max_streak=context.user_data.get("max_streak", 0),
        incorrect=incorrect,
        language=language,
    )
    temporary_path = ""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temporary:
            temporary_path = temporary.name
        document.save(temporary_path)
        with open(temporary_path, "rb") as stream:
            await context.bot.send_document(
                chat_id=update.effective_chat.id,
                document=stream,
                filename=safe_report_filename(test_name, start, end),
            )
    finally:
        if temporary_path and os.path.exists(temporary_path):
            os.unlink(temporary_path)
